from flask import Flask, request, jsonify
import pymysql
import jieba.analyse
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain_core.documents import Document
import os
import io
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
from langchain_ollama import OllamaLLM
from ollama._types import ResponseError
import time
import ollama  # 直接导入ollama用于诊断

app = Flask(__name__)


# 自定义嵌入模型类
class SentenceTransformerEmbeddings:
    def __init__(self, model_name):
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, texts):
        return self.model.encode(texts).tolist()

    def embed_query(self, text):
        return self.model.encode([text])[0].tolist()


# 连接 MySQL 数据库
def connect_to_mysql():
    connection = pymysql.connect(
        host='localhost',
        user='root',
        password='123456',
        database='linknote',
        cursorclass=pymysql.cursors.DictCursor
    )
    return connection


# 获取用户知识库文件
def get_user_knowledge_base(user_id):
    connection = connect_to_mysql()
    try:
        with connection.cursor() as cursor:
            sql = "SELECT file_content, file_type FROM knowledge_base WHERE user_id = %s"
            cursor.execute(sql, (user_id,))
            results = cursor.fetchall()
            return results
    finally:
        connection.close()


# 提取不同格式文件的文本内容
def extract_text_from_file(file_content, file_type):
    if not file_content:
        return ""
    if file_type == 'pdf':
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file_type == 'docx':
        doc = DocxDocument(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text
        return text
    elif file_type == 'md':
        return file_content.decode('utf-8')
    return ""


# 提取关键词
def extract_keywords(text, topK=5):
    keywords = jieba.analyse.extract_tags(text, topK=topK)
    return keywords


# 构建向量数据库并提取关键词
def build_vectorstore(file_data_list):
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)  # 调整分块大小
    docs = []
    file_keywords = []
    for file_data in file_data_list:
        file_content = file_data['file_content']
        file_type = file_data['file_type']
        text = extract_text_from_file(file_content, file_type)

        # 提取关键词
        keywords = extract_keywords(text)
        file_keywords.append(keywords)

        # 手动创建 Document 对象
        document = Document(page_content=text)
        documents = text_splitter.split_documents([document])
        docs.extend(documents)

    # 使用自定义的 SentenceTransformerEmbeddings 作为嵌入模型
    embeddings = SentenceTransformerEmbeddings('all-MiniLM-L6-v2')
    vectorstore = Chroma.from_documents(docs, embeddings)
    print(f"Number of documents in vectorstore: {len(docs)}")  # 打印文档数量
    return vectorstore, file_keywords


# 直接使用 Ollama 客户端进行问答
def direct_ollama_query(context, question, max_retries=3):
    client = ollama.Client(host="http://localhost:11434")

    # 构建完整提示
    prompt = f"""
    基于以下上下文回答问题。如果上下文中没有足够的信息，就说你不知道。

    上下文:
    {context}

    问题: {question}
    """

    # 使用重试机制
    for attempt in range(max_retries):
        try:
            # 使用您已经成功测试过的模型名称
            response = client.chat(
                model='deepseek-r1:latest',
                messages=[{'role': 'user', 'content': prompt}],
                stream=False
            )
            return response['message']['content']
        except Exception as e:
            print(f"尝试 {attempt + 1}/{max_retries} 失败: {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(5)  # 等待5秒再重试

    return "抱歉，我无法处理您的请求。请稍后再试。"


# 初始化本地 DeepSeek 模型并创建 RAG 链
def init_rag_chain(vectorstore):
    # 由于我们现在使用直接的ollama客户端，此函数主要返回retriever
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})  # 增加检索文档数量到3
    return retriever


@app.route('/ask', methods=['POST'])
def ask():
    data = request.get_json()
    question = data.get('question')
    user_id = data.get('user_id')

    # 获取用户知识库文件
    file_data_list = get_user_knowledge_base(user_id)

    # 构建向量数据库并提取关键词
    vectorstore, file_keywords = build_vectorstore(file_data_list)

    # 获取检索器
    retriever = init_rag_chain(vectorstore)

    # 检索相关文档
    relevant_docs = retriever.get_relevant_documents(question)

    # 提取文档内容作为上下文
    context = "\n".join([doc.page_content for doc in relevant_docs])

    # 直接使用Ollama客户端进行问答
    result = direct_ollama_query(context, question)

    response = {
        'answer': result,
        'file_keywords': file_keywords
    }
    return jsonify(response)


# 添加健康检查路由，用于测试Ollama连接
@app.route('/health', methods=['GET'])
def health_check():
    try:
        client = ollama.Client(host="http://localhost:11434")
        models = client.list()
        return jsonify({
            'status': 'ok',
            'models': [model.model for model in models.models]
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)